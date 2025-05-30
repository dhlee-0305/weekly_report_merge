import datetime
import os

from docx import Document

from weekly_env import *
from logger import *
from elapsed import *

from weekly_load import *
from status_project import *
from status_manpower import *
from status_client import *
from status_work import *

@elapsed
def main():

    # 작업 폴더 생성
    checkEnv()

    config = loadConfig(CONFIG_FILE)
    log = getLogger('weekly_merge')
    log.info("==> Weekly Report Merge Start!")

    weeklyFileList = getWeeklyList(config['PATH']['REPORT_PATH'])
    weeklyMergeDoc = Document(config['PATH']['TEMPLATE_REPORT_FILE'])

    # 상단 탑 영역 : 보고일 업데이트
    loadWeeklyTop(weeklyMergeDoc.tables[0], config['ENV']['REFERENCE_DAYOFWEEK'])

    for index in range(len(weeklyFileList)):
        log.info("["+str(index+1)+"] 주간보고 파일 : "+weeklyFileList[index])
        teamWeeklyDoc = Document(config['PATH']['REPORT_PATH']+weeklyFileList[index])
            
        # 1. 프로젝트 진행 현황
        if config.getboolean('MERGE_SECTION', 'PROJECT_STATUS') :
            projectStatusTable = teamWeeklyDoc.tables[1]
            loadProjectStatus(projectStatusTable, weeklyMergeDoc.tables[1])

        # 2. 인력 운용 현황
        if config.getboolean('MERGE_SECTION', 'MANPOWER_STATUS'):
            manpowerStatusTable = teamWeeklyDoc.tables[2]
            loadManpowerStatus(manpowerStatusTable, weeklyMergeDoc.tables[2])

        # 3. 거래처 영업/동향 정보
        if config.getboolean('MERGE_SECTION', 'CLIENT_STATUS'):
            clientStatusTable = teamWeeklyDoc.tables[3]
            loadClientStatus(clientStatusTable, weeklyMergeDoc.tables[3])

        # 4. 주요 업무 사항
        if config.getboolean('MERGE_SECTION', 'WORK_DESCRIPTION'):
            loadWork(teamWeeklyDoc, weeklyMergeDoc)

    try:
        weeklyMergeDocCreatePath = config['PATH']['OUTPUT_PATH'] + getWeeklyFilePrefix(config['ENV']['REFERENCE_DAYOFWEEK']) + '_IT서비스부문_주간보고.docx'
        weeklyMergeDoc.save(weeklyMergeDocCreatePath)
        log.info("==> Weekly Merge Complete -> " + weeklyMergeDocCreatePath)
    except:
        log.info("==> Weekly Merge Write Error, close opened file")

if __name__ == '__main__':
    main()