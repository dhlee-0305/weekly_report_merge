from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING

from weekly_env import *
from paragraph_style import *
from logger import *
from elapsed import *
from weekly_load import *

log = getLogger('status_work')

# 4. 주요 업무 사항
def loadWork(teamWeeklyDoc, weeklyMergeDoc):
    """ 
    주요 업무 사항을 읽어와 취합 문서에 추가한다.
    
    :param teamWeeklyDoc: 소스(팀) 문서
    :param weeklyMergeDoc: 취합 문서
    :return: None
    """
    # 팀이름 가져오기
    teamName = loadTeamName(teamWeeklyDoc.tables[0])

    log.debug('4. 주요 업무 사항')

    find_work = False
    for i, paragraph in enumerate(teamWeeklyDoc.paragraphs):
        if (paragraph.text == "주요 업무 사항"):
            find_work = True
            insertParagraph = weeklyMergeDoc.add_paragraph(teamName)
            insertParagraph.alignment = WD_TABLE_ALIGNMENT.LEFT
            setFontSizeBold(insertParagraph, 11, True)
        elif find_work == True:
            if(len(paragraph.text) > 0):
                insertParagraph = weeklyMergeDoc.add_paragraph(paragraph.text)
                insertParagraph.style = paragraph.style