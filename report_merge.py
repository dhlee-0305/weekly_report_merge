import datetime
import os

from docx import Document

from report_env import *
from report_load import *
from logger import *
from elapsed import *

@elapsed
def main():

    config = loadConfig(CONFIG_FILE)
    log = getLogger('report_merge')
    log.info("==> Weekly Report Merge Start!")

    report_file_list = getReportList(config['PATH']['REPORT_PATH'])
    dstDoc = Document(config['PATH']['TEMPLATE_REPORT_FILE'])

    # 상단 탑 영역 : 보고일 업데이트
    loadTop(dstDoc.tables[0], config['ENV']['REFERENCE_DAYOFWEEK'])

    for index in range(len(report_file_list)):
        log.info("["+str(index+1)+"] 주간보고 파일 : "+report_file_list[index])
        srcDoc = Document(config['PATH']['REPORT_PATH']+report_file_list[index])
            
        # 0. 상단 탑 영역 : 팀 이름 확인
        topTable = srcDoc.tables[0]
        teamName = loadTeamName(topTable)

        # 1. 프로젝트 진행 현황
        if config.getboolean('MERGE_SECTION', 'PROJECT_STATUS') :
            projectStatusTable = srcDoc.tables[1]
            loadProjectStatus(projectStatusTable, dstDoc.tables[1])

        # 2. 인력 운용 현황
        if config.getboolean('MERGE_SECTION', 'MANPOWER_STATUS'):
            manpowerStatusTable = srcDoc.tables[2]
            loadManpowerStatus(manpowerStatusTable, dstDoc.tables[2])

        # 3. 거래처 영업/동향 정보
        if config.getboolean('MERGE_SECTION', 'CLIENT_STATUS'):
            clientStatusTable = srcDoc.tables[3]
            loadClientStatus(clientStatusTable, dstDoc.tables[3])

        # 4. 주요 업무 사항
        if config.getboolean('MERGE_SECTION', 'WORK_DESCRIPTION'):
            loadWork(srcDoc, dstDoc, teamName)

    try:
        reportCreatePath = config['PATH']['OUTPUT_PATH'] + getReportFilePrefix(config['ENV']['REFERENCE_DAYOFWEEK']) + '_IT서비스부문_주간보고.docx'
        dstDoc.save(reportCreatePath)
        log.info("==> Report Merge Complete -> " + reportCreatePath)
    except:
        log.info("==> Report Merge Write Error, close opened file")

if __name__ == '__main__':
    main()