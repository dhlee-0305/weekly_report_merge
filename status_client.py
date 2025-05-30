from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING

from weekly_env import *
from paragraph_style import *
from logger import *
from weekly_load import *

log = getLogger('status_client')

# 3. 거래처 영업/동향 정보
def loadClientStatus(teamWeeklyDoc, weeklyMergeDoc):
    """ 
    거래처 영업/동향을 읽어와 취합 문서에 추가한다.
    
    :param teamWeeklyDoc: 소스(팀) 문서
    :param weeklyMergeDoc: 취합 문서
    :return: None
    """    
    log.debug('3. 거래처 영업/동향 정보')

    rowSize, colSize = getColRowSize(teamWeeklyDoc)
    
    if colSize != CLIENT_STATUS_COL_SIZE:
        log.error("loadClientStatus 거래처 영업/동향 정보 column size invalid: "+ str(colSize))
        exit()


    for row in range(rowSize-1):
        if (len(teamWeeklyDoc.rows[row+1].cells[2].text) > 3) :
            rowCells = weeklyMergeDoc.add_row().cells
        else:
            # 공백 행 제외
            continue
        
        for col in range(colSize):
            if col == 2: # (2)주요 정보
                customerIssueStrArray = teamWeeklyDoc.rows[row+1].cells[col].text.strip().split('\n')
                isTitle = True
                for customerStrIssueStr in customerIssueStrArray:
                    if len(customerStrIssueStr) > 1:
                        if isTitle :
                            insertParagraph = rowCells[col].paragraphs[0]
                            insertParagraph.add_run(customerStrIssueStr)
                        else:
                            if(customerStrIssueStr.startswith('-')):
                                insertParagraph = rowCells[col].add_paragraph(' ' + customerStrIssueStr) # 서식 무시하고 '-' 추가한 경우 제외를 위함
                            else:
                                insertParagraph = rowCells[col].add_paragraph(' - ' + customerStrIssueStr)
                        
                        setStyleCustomerIssue(insertParagraph, isTitle)
                        isTitle = False

            else: # (0)구분, (1)고객사/부서, (3)비고
                rowCells[col].text = teamWeeklyDoc.rows[row+1].cells[col].text

                rowCells[col].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                setFontSizeBold(rowCells[col].paragraphs[0], 9, False)

def setStyleCustomerIssue(insertParagraph, isTitle):
    """
    거래처 영업/동향 정보의 주요 정보 내용 스타일 설정
    :param insertParagraph: 단락
    :param isTitle: 제목 여부
    :return: 없음
    """
    
    insertParagraph.alignment = WD_TABLE_ALIGNMENT.LEFT
    paragraph_format = insertParagraph.paragraph_format
    paragraph_format.line_spacing = Pt(12)
    paragraph_format.space_before = Pt(5)
    
    run_obj = insertParagraph.runs
    run = run_obj[0]
    font = run.font
    font.size = Pt(9)

    if isTitle:
        font.bold = True
        isTitle = False
    else:
        font.bold = False